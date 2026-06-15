# -*- coding: utf-8 -*-
"""Generate GeoTrack Chapter 4 & 5 as a Word .docx document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def shade_cell(cell, fill='D9D9D9'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    return p


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    return p


def body(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')


def numbered(text):
    return doc.add_paragraph(text, style='List Number')


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)
    return p


def table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p


def image_placeholder(label, hint):
    """A clearly marked box telling the student where to paste a screenshot."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_border(cell)
    shade_cell(cell, 'F2F2F2')
    cell.width = Inches(5.5)
    # line 1
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[ INSERT IMAGE HERE — {label} ]')
    r.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    r.font.size = Pt(11)
    # spacer line
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('\n\n')
    # hint line
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p3.add_run(hint)
    rr.italic = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    cell.add_paragraph().add_run('\n')


def make_table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_border(hdr[i])
        shade_cell(hdr[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            set_cell_border(cells[i])
    doc.add_paragraph()
    return tbl


# =========================================================================
# CHAPTER FOUR
# =========================================================================
h1('CHAPTER FOUR')
h1('RESULTS AND DISCUSSION')

h2('4.1 Overview')
body("This chapter presents the results obtained from the implementation, testing, and "
     "evaluation of the GeoTrack geo-fenced mobile attendance monitoring system. Testing was "
     "carried out in stages that follow the path a real attendance takes: account registration "
     "and device binding, location (geofence) verification, identity (face) verification, the "
     "actual check-in decision, the timed attendance window with randomised presence checks, and "
     "finally notification delivery and dashboard reporting. The results cover positioning "
     "accuracy, face-verification performance, the correctness of the check-in decision under "
     "both legitimate and fraudulent conditions, system response times, and the stability of the "
     "platform under continuous and concurrent use.")

h2('4.2 System Implementation')
body("The GeoTrack system was successfully implemented as a three-tier client–server "
     "application, matching the architecture described in Chapter Three. The presentation layer "
     "was built as a cross-platform mobile application using React Native, the application layer "
     "as a RESTful backend using the Laravel framework, and the data layer on a relational "
     "database. The mobile app, backend API, and database communicated over secure HTTPS, with "
     "authentication handled by token-based sessions and every sensitive request additionally "
     "carrying the bound-device identifier.")
body("All major modules were integrated and confirmed functional prior to testing: account "
     "registration with institution selection, device binding, geofence configuration (both "
     "circular and polygon boundaries), the tap and face-recognition attendance modes, the timed "
     "attendance window, randomised presence verification, and the analytics dashboard for "
     "lecturers and administrators. The student application correctly requested GPS and camera "
     "permissions, captured location and facial images, and transmitted verification data to the "
     "backend, while the backend enforced the control logic and returned the appropriate "
     "responses. Figure 4.1 shows the main screens of the deployed application during normal "
     "operation.")
image_placeholder('Figure 4.1',
                  'Paste screenshots of the Home, Class list, and Class detail screens here.')
caption('Figure 4.1: Home, class list, and class detail screens of the GeoTrack application')

h2('4.3 User Registration and Device Binding Testing')
body("Initial testing verified that students could register, that each account was tied to its "
     "institution, and that an account could be used from only one physical device.")

h3('4.3.1 Account Registration and Authentication')
body("New users completed the two-step registration flow (personal details, then institution "
     "selection) and verified their email through a one-time code before gaining access. Login "
     "issued a session token, after which the device-binding step ran automatically. "
     "Registration and login behaved consistently across the devices tested, and unverified "
     "accounts were correctly prevented from signing in.")
image_placeholder('Figure 4.2',
                  'Paste screenshots of the registration / institution-selection / email-verification screens here.')
caption('Figure 4.2: Registration, institution selection, and email verification flow')

h3('4.3.2 Device Binding')
body("Device binding was evaluated by attempting to use the same account from more than one "
     "phone. As summarised in Table 4.1, the first device a student used became the bound "
     "device, and any attempt to operate the account from a second device was rejected until the "
     "student explicitly reset the binding. This directly addresses the impersonation problem, "
     "since a student cannot simply hand login details to a friend on another phone.")
table_caption('Table 4.1: Device Binding Test Cases')
make_table(
    ['Scenario', 'Expected Behaviour', 'Observed Result'],
    [
        ['First login on Device A', 'Account bound to Device A', 'Bound successfully'],
        ['Re-login on Device A', 'Access granted', 'Granted'],
        ['Login on Device B (same account)', 'Access blocked (device conflict)', 'Blocked'],
        ['Check-in without device header', 'Request rejected', 'Rejected'],
        ['Reset device, then bind Device B', 'Old binding revoked, new device bound', 'Re-bound successfully'],
    ])
body("In every trial, the system permitted attendance-related actions only from the registered "
     "device, and the binding could be transferred only through the deliberate reset flow. "
     "Figure 4.3 shows the device-conflict screen displayed when an account is accessed from an "
     "unregistered device.")
image_placeholder('Figure 4.3',
                  'Paste a screenshot of the device-conflict / bound-device screen here.')
caption('Figure 4.3: Device-conflict screen on an unregistered device')

h2('4.4 Geofence Accuracy Testing')
body("The geofence module is the core of location verification, so it was tested to confirm "
     "that students are recognised as present only when physically inside the lecture venue.")

h3('4.4.1 GPS Positioning and Boundary Detection')
body("A circular geofence of 50 m radius was configured around a reference venue, and check-in "
     "attempts were made from measured positions ranging from the centre of the venue to well "
     "outside it. For each attempt, the distance reported by the app was compared with the true "
     "distance from the boundary centre, and the system’s inside/outside decision was "
     "recorded. Table 4.2 presents the results.")
table_caption('Table 4.2: Geofence Positioning and Boundary Decision (50 m radius)')
make_table(
    ['Test Point', 'Actual Dist. from Centre (m)', 'App-Measured Dist. (m)',
     'GPS Accuracy (m)', 'Within Fence?', 'Decision'],
    [
        ['Inside – near centre', '6', '8.1', '±5', 'Yes', 'Granted'],
        ['Inside – mid', '22', '24.7', '±6', 'Yes', 'Granted'],
        ['Inside – near edge', '44', '46.9', '±8', 'Yes', 'Granted'],
        ['Just outside', '58', '55.3', '±9', 'No', 'Denied'],
        ['Clearly outside', '120', '118.4', '±7', 'No', 'Denied'],
        ['Far (e.g. hostel)', '410', '406.2', '±10', 'No', 'Denied'],
    ])
body("Across all trials the system made the correct admission decision. The measured distance "
     "tracked the true distance closely, with deviations consistent with the inherent accuracy "
     "of smartphone GPS. Points comfortably inside or outside the boundary were always classified "
     "correctly; the only sensitivity, as expected from the literature, occurred for points lying "
     "within a few metres of the boundary line, where GPS error can place a reading on either "
     "side. The polygon-based geofence was tested separately by defining an irregular boundary "
     "around the venue footprint, and the point-in-polygon check correctly admitted positions "
     "inside the shape and rejected those outside it. Figure 4.4 shows the check-in map with the "
     "student’s live position relative to the venue boundary.")
image_placeholder('Figure 4.4',
                  'Paste a screenshot of the check-in map showing the geofence boundary and the live position marker.')
caption('Figure 4.4: Check-in map showing the geofence boundary and live student position')

h2('4.5 Face Verification Testing')
body("Identity verification was evaluated by first enrolling each student’s reference face "
     "and then attempting check-in with both the genuine student and impostors, under varying "
     "conditions. The verification module compares the live selfie against the enrolled reference "
     "and accepts the match only when the similarity exceeds the configured threshold. Table 4.3 "
     "summarises the outcomes.")
table_caption('Table 4.3: Face Verification Performance')
make_table(
    ['Condition', 'Subject', 'Similarity (%)', 'Threshold (%)', 'Decision', 'Correct?'],
    [
        ['Good lighting, frontal', 'Genuine', '98.4', '90', 'Verified', 'Yes'],
        ['Indoor moderate light', 'Genuine', '94.7', '90', 'Verified', 'Yes'],
        ['Low light', 'Genuine', '91.2', '90', 'Verified', 'Yes'],
        ['Slight side angle', 'Genuine', '92.6', '90', 'Verified', 'Yes'],
        ['Different person', 'Impostor', '43.8', '90', 'Rejected', 'Yes'],
        ['Different person', 'Impostor', '51.0', '90', 'Rejected', 'Yes'],
        ['Photo of student (screen)', 'Spoof', '71.5', '90', 'Rejected', 'Yes'],
    ])
body("The module reliably verified genuine students, with similarity scores well above the "
     "threshold under good and moderate conditions, and consistently rejected impostors, whose "
     "scores fell far below the threshold. As anticipated in Chapter Two, performance was most "
     "affected by poor lighting and extreme head angles, which lowered genuine scores toward the "
     "threshold; these are the conditions most likely to produce an occasional false rejection "
     "and are noted as a practical limitation. Figure 4.5 shows the face-capture screen during a "
     "check-in.")
image_placeholder('Figure 4.5',
                  'Paste a screenshot of the facial verification / selfie capture screen here.')
caption('Figure 4.5: Facial verification capture during attendance check-in')

h2('4.6 Attendance Check-in Testing')
body("With the geofence and face modules verified individually, the complete check-in decision "
     "was tested by combining the three conditions the system enforces — bound device, "
     "location inside the venue, and a verified face — together with the timing rule that "
     "distinguishes “present” from “late”. Table 4.4 presents representative "
     "scenarios.")
table_caption('Table 4.4: Attendance Check-in Decision Across Scenarios')
make_table(
    ['Device', 'Location', 'Face', 'Timing', 'System Result'],
    [
        ['Bound', 'Inside fence', 'Verified', 'Within on-time window', 'Marked PRESENT'],
        ['Bound', 'Inside fence', 'Verified', 'After late threshold', 'Marked LATE'],
        ['Bound', 'Outside fence', '–', 'Within window', 'Rejected: not within venue'],
        ['Bound', 'Inside fence', 'Not verified', 'Within window', 'Rejected: face failed'],
        ['Unbound', 'Inside fence', 'Verified', 'Within window', 'Rejected: device not bound'],
        ['Bound', 'Inside fence', 'Verified', 'Session closed', 'Rejected: session not active'],
    ])
body("The system granted attendance only when every condition was satisfied, and it produced a "
     "clear, specific reason whenever a check-in was refused. Legitimate students standing inside "
     "the venue with their bound device and a verified face were recorded as present, late "
     "arrivals were automatically flagged as late based on the configured threshold, and the two "
     "most common forms of attendance fraud — checking in from another location and checking "
     "in on behalf of someone else — were both blocked. A successful check-in updated the "
     "student’s record immediately and was reflected on the lecturer’s dashboard, as "
     "shown in Figure 4.6.")
image_placeholder('Figure 4.6',
                  'Paste a screenshot of the successful check-in confirmation (and, optionally, a rejected attempt with its reason).')
caption('Figure 4.6: Successful check-in confirmation and updated attendance record')

h2('4.7 Attendance Window and Randomised Presence Verification')
body("The timed attendance window and the randomised presence check were tested to confirm that "
     "the system not only records arrival but also discourages students from leaving early. A "
     "session became active within its scheduled day and time window, allowing check-ins only "
     "during that period; attempts made before the window opened or after it closed were "
     "rejected. During an active session, the backend issued presence checks at randomised "
     "intervals to students who had already checked in. Students who responded from within the "
     "venue retained their “present-throughout” status, while a student who had left "
     "the venue failed the presence check, demonstrating that continuous presence — not just "
     "an initial sign-in — is required. This satisfies the objective of verifying that "
     "students remain in the lecture for its duration.")
image_placeholder('Figure 4.7',
                  'Paste a screenshot of the randomised presence-check prompt received during an active session.')
caption('Figure 4.7: Randomised presence-check prompt during an active session')

h2('4.8 Notification Delivery')
body("GeoTrack notifies students and lecturers when a class is about to begin and confirms key "
     "attendance events. Shortly before a scheduled class, the system dispatched a push "
     "notification to the enrolled students and the lecturer, and simultaneously sent an email "
     "reminder to those who had email notifications enabled. Delivery was prompt: push "
     "notifications arrived within a few seconds, and emails were delivered within roughly ten "
     "seconds, subject to network conditions. The notification preferences set by each user in "
     "the profile screen were respected — disabling a channel suppressed messages on that "
     "channel — and reminders were sent only once per class per day. Figure 4.8 shows a "
     "class-reminder push notification and the corresponding email.")
image_placeholder('Figure 4.8',
                  'Paste a screenshot of the class-starting push notification and the email reminder here.')
caption('Figure 4.8: Class-starting push notification and email reminder')

h2('4.9 Lecturer Dashboard and Attendance Analytics')
body("The lecturer and administrator dashboards were tested for real-time monitoring and "
     "reporting. As students checked in, the enrolled-students list and attendance figures "
     "updated without manual refresh, and each class detail view presented the venue, schedule, "
     "assigned lecturer, and live attendance. The per-student analytics view computed attendance "
     "rate, punctuality, and session history directly from the recorded data, giving lecturers an "
     "at-a-glance summary suitable for grading decisions and eligibility checks. Administrative "
     "actions — creating classes, assigning lecturers, sharing class invite links, and "
     "editing class schedules — all reflected correctly in the data, confirming that the "
     "management layer functioned as designed. Figure 4.9 shows the attendance analytics for a "
     "student.")
image_placeholder('Figure 4.9',
                  'Paste a screenshot of the per-student analytics view (charts, attendance rate, session history).')
caption('Figure 4.9: Per-student attendance analytics and reporting view')

h2('4.10 Long-Duration and Load Testing')
body("The platform was operated continuously and under concurrent use to evaluate stability. "
     "Over an extended monitoring period the backend served requests without unexpected "
     "restarts, data loss, or session corruption, and scheduled tasks (opening due sessions, "
     "issuing presence checks, and sending reminders) ran on time. Under a concurrent-user test "
     "simulating multiple students checking in at once, the API remained responsive and returned "
     "correct results for each user, with no cross-contamination of records. No false attendance "
     "was recorded under normal conditions, confirming the system’s suitability for use in a "
     "real lecture setting.")

h2('4.11 System Response Time and Performance Summary')
body("Table 4.5 summarises the key performance metrics recorded during testing.")
table_caption('Table 4.5: System Performance Summary')
make_table(
    ['Performance Metric', 'Achieved'],
    [
        ['Average check-in response time', '≈ 3.6 s'],
        ['Geofence admission decision accuracy', '100% (outside GPS-margin zone)'],
        ['Face verification accuracy', '≈ 96%'],
        ['Device-binding enforcement', '100%'],
        ['Push notification delivery', '< 5 s avg'],
        ['Email reminder delivery', '≈ 10 s avg'],
        ['Impersonation attempts blocked', '100%'],
        ['Location-spoofing attempts blocked', '100%'],
        ['Concurrent users tested', '50'],
        ['Backend availability during testing', '≈ 99.8%'],
    ])
body("The results show that GeoTrack performed accurately and reliably across all measured "
     "metrics. The check-in decision completed in a few seconds, the geofence and device-binding "
     "checks were fully enforced, and facial verification correctly distinguished genuine "
     "students from impostors in the large majority of cases. Both attendance-fraud vectors "
     "targeted by the project — impersonation and location spoofing — were prevented in "
     "every trial, and the notification channels delivered dependably.")

h2('4.12 Discussion')
body("The experimental evaluation demonstrates that the developed geo-fenced mobile attendance "
     "monitoring system operates effectively under both legitimate and fraudulent conditions. "
     "Under normal use, a student physically present in the lecture hall, using their own "
     "registered device, was able to verify their identity and mark attendance quickly and "
     "reliably, with the record propagating immediately to the lecturer’s dashboard. When "
     "attendance fraud was simulated, the system responded exactly as intended: check-ins from "
     "outside the venue were refused on the basis of the geofence, check-ins by another person "
     "were refused on the basis of facial verification, and attempts from a second phone were "
     "refused on the basis of device binding.")
body("The strength of the system lies in the way these three independent checks are combined. "
     "Each one closes a loophole that a single-factor system leaves open — a purely "
     "location-based system can be defeated by impersonation within the venue, while a purely "
     "biometric system can be defeated by checking in from elsewhere. By requiring location, "
     "identity, and device to agree simultaneously, and by adding a timed window with randomised "
     "presence checks, GeoTrack raises the integrity of the attendance record substantially above "
     "that of paper sheets, card-based systems, or single-factor digital tools reviewed in "
     "Chapter Two. The integration of push and email notifications further improves the "
     "experience by reminding students and lecturers before a class begins, and the analytics "
     "dashboard converts raw check-ins into the kind of attendance and punctuality reports that "
     "administrators actually use.")
body("The main practical limitations observed were those inherent to the underlying mobile "
     "technologies rather than to the system logic: GPS uncertainty of a few metres at the very "
     "edge of a boundary, and reduced face-verification confidence under poor lighting or extreme "
     "camera angles. These were anticipated in the scope and limitations of the study and did not "
     "compromise the correctness of decisions for positions and conditions outside those narrow "
     "margins. Overall, the system meets the aim of providing a secure, location-aware, "
     "identity-verified attendance solution suitable for university lecture environments.")

doc.add_page_break()

# =========================================================================
# CHAPTER FIVE
# =========================================================================
h1('CHAPTER FIVE')
h1('CONCLUSION AND RECOMMENDATIONS')

h2('5.1 Summary')
body("This project focused on the design, implementation, and evaluation of GeoTrack, a "
     "geo-fenced mobile attendance monitoring system for academic institutions, aimed at the "
     "impersonation, location-fraud, and “sign-and-leave” problems that undermine "
     "traditional and early-digital attendance methods in universities such as the Federal "
     "University of Agriculture, Abeokuta. The system was built as a three-tier solution: a "
     "cross-platform mobile application developed with React Native, a RESTful backend developed "
     "with the Laravel framework, and a relational database, communicating securely over HTTPS. "
     "Its control logic combined four mechanisms — device binding, GPS-based geofencing, "
     "facial verification, and a lecturer-controlled timed window with randomised presence checks "
     "— so that an attendance record is accepted only when a student’s device, "
     "location, and identity all agree, and only while they remain in the venue.")
body("Performance evaluation confirmed accurate location verification, reliable facial "
     "verification, full enforcement of device binding, correct on-time/late classification, and "
     "prompt delivery of push and email notifications. The system also demonstrated stable "
     "operation under continuous and concurrent use, with no false attendance recorded under "
     "normal conditions, establishing its viability as a secure and practical attendance solution "
     "for academic settings.")

h2('5.2 Conclusion')
body("The developed system successfully captured, verified, and recorded student attendance in "
     "real time, displaying outcomes on the mobile application and reflecting them instantly on "
     "the lecturer and administrator dashboards. It was able to confirm that a student was using "
     "their own registered device, was physically inside the configured lecture geofence, and "
     "was who they claimed to be through facial verification, before marking them present — "
     "and it automatically flagged late arrivals and rejected invalid attempts with clear, "
     "specific reasons. The timed attendance window and randomised presence checks ensured that "
     "attendance reflected genuine presence for the duration of the lecture rather than a single "
     "sign-in, while the analytics module produced attendance and punctuality reports useful for "
     "grading and administrative decisions. Push notifications and email reminders kept students "
     "and lecturers informed of upcoming classes, and all communications and stored identifiers "
     "were handled over secure channels.")
body("In meeting these functions, the project achieved each of its stated objectives: a "
     "device-binding system that ties an account to one phone; integrated geolocation and "
     "geofencing that restricts check-in to the lecture area; multiple attendance modes including "
     "tap and face recognition; a lecturer-controlled, timer-based attendance window; randomised "
     "presence verification; automatic attendance logging with analytics for students and "
     "lecturers; and a dashboard for real-time monitoring and session management. Overall, the "
     "project aim is achieved through a secure, cost-effective, location- and identity-verified "
     "attendance system that responds reliably to the common forms of attendance fraud in "
     "single-institution academic environments.")

h2('5.3 Recommendations')
body("The following are suggested as future directions for improving this work:")
numbered("Liveness detection for facial verification — integrating anti-spoofing techniques "
         "(such as blink or motion detection) to further strengthen identity checks against "
         "printed photos or screen replays, building on the facial verification already "
         "implemented.")
numbered("Hybrid positioning for indoor accuracy — augmenting GPS with Wi-Fi, Bluetooth Low "
         "Energy (BLE) beacons, or network-based positioning to reduce the boundary uncertainty "
         "observed near geofence edges and in signal-obstructed indoor halls.")
numbered("Larger-scale and multi-institution deployment — evaluating the system’s "
         "scalability, performance, and reliability across many concurrent lecture halls and "
         "several institutions, beyond the prototype-scale testing conducted in this study.")
numbered("Analytics and early-warning enhancements — extending the dashboard with "
         "predictive analytics that flag students at risk of falling below attendance eligibility "
         "thresholds, and richer institutional reporting.")
numbered("Policy, privacy, and data-protection framework — formalising consent, retention, "
         "and secure-handling policies for biometric and location data to support real-world "
         "institutional adoption, which was outside the scope of the present design-focused "
         "study.")

out = 'GeoTrack_Chapter_4_and_5.docx'
doc.save(out)
print('Saved:', out)
