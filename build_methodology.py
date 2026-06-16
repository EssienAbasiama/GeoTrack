# -*- coding: utf-8 -*-
"""Generate GeoTrack Chapter Three (Methodology) with engineering formulas as .docx.

Formulas are taken directly from the implemented source code:
  - GeofenceService.php   -> Haversine distance, circular & polygon geofence
  - FaceMatchService.php   -> grayscale luminance, average hash, Hamming distance,
                              confidence, AWS Rekognition similarity
  - AttendanceController.php -> lateness, present/late classification, composite gate
  - PresenceCheckService.php -> randomised presence checks, present-throughout
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

# Page margins (typical thesis: 1.5" left, 1" others)
sec = doc.sections[0]
sec.left_margin = Inches(1.5)
sec.right_margin = Inches(1.0)
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tcPr.append(borders)


def shade_cell(cell, fill='D9D9D9'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    return p


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')


def numbered(text):
    return doc.add_paragraph(text, style='List Number')


def equation(parts, number):
    """Render a centred equation line with a right-aligned equation number.

    `parts` is a list of (text, italic) tuples so variables can be italicised
    in true mathematical style; `number` is e.g. '3.1'.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Tab stops: centre for the equation, right for the number.
    pf.tab_stops.add_tab_stop(Inches(2.6), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run('\t')
    for txt, ital in parts:
        r = p.add_run(txt)
        r.italic = ital
        r.font.size = Pt(12)
    p.add_run('\t')
    rn = p.add_run('(' + number + ')')
    rn.font.size = Pt(12)
    return p


def where(lines):
    """A 'where:' definition block under an equation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('where:')
    r.italic = True
    for var, desc in lines:
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Inches(0.5)
        lp.paragraph_format.space_after = Pt(0)
        rv = lp.add_run(var)
        rv.italic = True
        lp.add_run('  ' + desc)


def figure_placeholder(label, caption_text):
    """A bordered box standing in for a figure to be inserted later."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER if hasattr(WD_ALIGN_PARAGRAPH, 'CENTER') else None
    cell = tbl.rows[0].cells[0]
    set_cell_border(cell)
    cell.width = Inches(5.5)
    ph = cell.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run('[ INSERT ' + label + ' HERE ]')
    r.bold = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for _ in range(3):
        cell.add_paragraph()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(caption_text)
    rc.bold = True
    rc.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(12)


def simple_table(headers, rows, caption_text):
    cap = doc.add_paragraph()
    rc = cap.add_run(caption_text)
    rc.bold = True
    rc.font.size = Pt(11)
    cap.paragraph_format.space_before = Pt(8)
    tbl = doc.add_table(rows=1, cols=len(headers))
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_border(hdr[i])
        shade_cell(hdr[i])
        hp = hdr[i].paragraphs[0]
        rr = hp.add_run(htext)
        rr.bold = True
        rr.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_border(cells[i])
            cp = cells[i].paragraphs[0]
            rr = cp.add_run(str(val))
            rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ============================================================================
# CHAPTER THREE
# ============================================================================
h1('CHAPTER THREE')
h1('MATERIALS AND METHODS')

# ---- 3.1 Introduction ----
h2('3.1  Introduction')
body(
    'This chapter presents the materials and methods adopted in the design and '
    'implementation of GeoTrack, a geo-fenced mobile attendance monitoring system '
    'for academic institutions. It describes the research design, the hardware and '
    'software materials employed, the overall system architecture, and the '
    'mathematical models and algorithms that govern each functional subsystem. '
    'Particular emphasis is placed on the engineering formulations used for '
    'satellite-based positioning and geofencing, biometric (face) verification, '
    'device binding, attendance-time classification, and continuous presence '
    'verification, since these constitute the analytical core of the system.'
)

# ---- 3.2 Research Design ----
h2('3.2  Research Design')
body(
    'A Design and Development Research (DDR) approach was adopted, providing a '
    'systematic progression through problem analysis, system modelling, '
    'implementation, and evaluation. The system was realised as a three-tier '
    'client–server architecture comprising a presentation layer, an '
    'application (logic) layer, and a data layer, chosen for modularity, '
    'maintainability, and horizontal scalability. The presentation layer is a '
    'cross-platform React Native mobile application; the application layer is a '
    'Laravel RESTful backend exposing the system’s business logic over secure '
    'HTTPS; and the data layer is a relational database that persists users, '
    'courses, geofences, sessions, and attendance records. Embedded smartphone '
    'sensors (GPS receiver and front-facing camera) together with wireless '
    'connectivity provide the real-time data-acquisition channel for attendance '
    'verification.'
)

# ---- 3.3 Materials ----
h2('3.3  Materials')

h3('3.3.1  Hardware Components')
bullet('Smartphones (Android/iOS) equipped with GPS receivers, front-facing cameras, accelerometers, and wireless communication modules, serving as the data-acquisition and user-interface terminals.')
bullet('Cloud/server infrastructure (Linux–Ubuntu) for application logic, data processing, storage, and scheduled background tasks.')
bullet('Development workstations (laptops/desktops) for software development, simulation, and testing.')

h3('3.3.2  Software and Development Tools')
bullet('Mobile application: React Native with TypeScript/JavaScript (Expo) for cross-platform deployment.')
bullet('Backend and API: Laravel (PHP) implementing a RESTful API architecture with Sanctum token authentication.')
bullet('Database: relational SQL database (MySQL/SQLite) for normalised data storage and query optimisation.')
bullet('Cloud infrastructure: Nginx/Apache web server with SSL/TLS security and automated backups.')

h3('3.3.3  External APIs and Libraries')
bullet('Google Maps and Geolocation APIs for accurate positioning and map rendering.')
bullet('Geofencing logic (Haversine and ray-casting algorithms) to enforce spatial boundaries.')
bullet('Face-recognition services (AWS Rekognition deep-CNN matching, with a local perceptual-hash fallback) for biometric verification.')
bullet('Device-fingerprinting libraries for secure, unique device identification.')
bullet('Axios for secure client–server communication over HTTPS.')

# ---- 3.4 System Architecture ----
h2('3.4  System Architecture')
body(
    'The system follows the three-tier model illustrated in the block diagram '
    '(Figure 3.1). All inter-layer communication occurs over encrypted HTTPS '
    'channels; every sensitive request carries a bearer session token and a '
    'unique bound-device identifier (X-Device-UID) that the backend validates '
    'before any attendance operation is permitted. The application layer is '
    'decomposed into cooperating services — a Geofence Service, a Face-Match '
    'Service, a Scheduled-Session Service, and a Presence-Check Service — each '
    'encapsulating one verification responsibility.'
)
figure_placeholder('SYSTEM BLOCK DIAGRAM', 'Figure 3.1: Three-tier system block diagram of GeoTrack')
figure_placeholder('SYSTEM FLOWCHART', 'Figure 3.2: System workflow / attendance check-in flowchart')

# ---- 3.5 Subsystem design and mathematical models ----
h2('3.5  Subsystem Design and Mathematical Models')
body(
    'This section formalises the algorithms implemented in each verification '
    'subsystem. The equations below correspond directly to the deployed source '
    'code and define the decision rules the system applies at runtime.'
)

# 3.5.1 Geolocation & Geofencing
h3('3.5.1  Geolocation and Geofencing')
body(
    'A student’s instantaneous position is obtained from the smartphone GNSS '
    'receiver as a latitude–longitude pair, P = (φₚ, λₚ), '
    'reported with an estimated horizontal accuracy a (in metres). Each lecture '
    'venue is modelled either as a circular geofence defined by a centre '
    'C = (φ₀, λ₀) and radius r, or as an arbitrary polygon defined '
    'by an ordered set of vertices. Because positions lie on the Earth’s '
    'curved surface, planar distance is invalid; the great-circle (orthodromic) '
    'distance is therefore computed using the Haversine formulation.'
)
body('The angular terms are first converted to radians:')
equation([('φ₁ = lat₁ · π/180,   φ₂ = lat₂ · π/180', False)], '3.1')
equation([('Δφ = (lat₂ − lat₁) · π/180,   Δλ = (lng₂ − lng₁) · π/180', False)], '3.2')
body('The Haversine intermediate term a and angular distance c are then:')
equation([('a = sin²(Δφ/2) + cos φ₁ · cos φ₂ · sin²(Δλ/2)', False)], '3.3')
equation([('c = 2 · atan2(√a, √(1 − a))', False)], '3.4')
body('and the surface distance between the two points is:')
equation([('d = R · c', False)], '3.5')
where([
    ('R', '= 6,371,000 m, the mean radius of the Earth;'),
    ('d', '= the great-circle distance between the student and the venue centre (m);'),
    ('φ, λ', '= latitude and longitude in radians, respectively.'),
])
body(
    'For a circular geofence, the student is judged to be physically present '
    'within the venue if, and only if, the computed distance does not exceed the '
    'configured radius:'
)
equation([('Inside', True), (' = ', False), ('true', False), (',   if  ', False), ('d ≤ r', False)], '3.6')
body(
    'For a polygon venue, presence is determined by the ray-casting (even–odd) '
    'algorithm: a horizontal ray is projected from the test point P, and the '
    'number of polygon edges it crosses is counted. The point is inside when the '
    'crossing count is odd. For each edge (Vᵢ, Vⱼ) the crossing test is:'
)
equation([
    ('(yᵢ > yₚ) ≠ (yⱼ > yₚ)  ∧  xₚ < (xⱼ − xᵢ)(yₚ − yᵢ)/(yⱼ − yᵢ) + xᵢ', False),
], '3.7')
where([
    ('(x, y)', '= (longitude, latitude) of the polygon vertices and the test point;'),
    ('odd crossings', '⇒ point is inside the polygon (attendance location valid).'),
])
body(
    'To accommodate GNSS uncertainty, the reported accuracy a is retained with '
    'each record so that decisions for points lying within a few metres of the '
    'boundary can be audited. Only positions whose distance is genuinely within '
    'the radius are admitted, ensuring that remote or proxy check-ins outside the '
    'venue are rejected.'
)

# 3.5.2 Face verification
h3('3.5.2  Biometric Identity Verification (Face Recognition)')
body(
    'To prevent impersonation, the identity of the person checking in is verified '
    'against a previously enrolled reference face. Two complementary techniques '
    'were implemented: a production-grade deep-learning matcher (AWS Rekognition) '
    'and a lightweight on-device perceptual-hash fallback. Both reduce a captured '
    'selfie to a compact descriptor and compare it with the enrolled descriptor '
    'using a similarity threshold.'
)
body(
    '(a) Perceptual average-hash matcher. The captured image is down-sampled to '
    'an N × N grid (N = 16, giving 256 pixels) and converted to greyscale using '
    'the standard luminance transform:'
)
equation([('Y = 0.299R + 0.587G + 0.114B', False)], '3.8')
where([
    ('Y', '= greyscale luminance of a pixel;'),
    ('R, G, B', '= the red, green, and blue colour components (0–255).'),
])
body('The mean intensity of all pixels is computed as:')
equation([('μ = (1/N²) · Σᵢ Yᵢ', False)], '3.9')
body(
    'A 256-bit binary hash H is then formed by thresholding each pixel against '
    'the mean, so the hash encodes the coarse structure of the face:'
)
equation([('bᵢ = 1  if  Yᵢ ≥ μ;   bᵢ = 0  otherwise', False)], '3.10')
body(
    'At verification, the Hamming distance between the enrolled hash Hᵣ and the '
    'live hash Hₗ counts the number of differing bits:'
)
equation([('D(Hᵣ, Hₗ) = Σᵢ₌₁²⁵⁶ (bᵣ,ᵢ ⊕ bₗ,ᵢ)', False)], '3.11')
body('This distance is mapped to a percentage confidence score:')
equation([('Confidence = 100 · (1 − D/256)  %', False)], '3.12')
body('and the match decision applies an acceptance threshold on the distance:')
equation([('Match', True), (' = ', False), ('true', False), (',   if  ', False), ('D ≤ D', False), ('max', False)], '3.13')
where([
    ('⊕', '= the bitwise exclusive-OR (XOR) operator;'),
    ('D', '= the Hamming distance between the two 256-bit hashes;'),
    ('D', 'max = the calibrated acceptance distance (configurable per deployment).'),
])
body(
    '(b) Deep-CNN matcher (AWS Rekognition). For production use, each face is '
    'encoded by a deep convolutional neural network into a high-dimensional '
    'embedding (feature) vector f ∈ ℝⁿ. The similarity between the enrolled '
    'embedding fᵣ and a live embedding fₗ is evaluated by cosine similarity:'
)
equation([('S(fᵣ, fₗ) = (fᵣ · fₗ) / (‖fᵣ‖ ‖fₗ‖)', False)], '3.14')
body(
    'which the service reports as a similarity percentage. A face is accepted '
    'only when the similarity meets or exceeds the configured threshold T (with '
    'T = 90% for production matching and T = 70% adopted for the demonstration '
    'environment):'
)
equation([('Verified', True), (' = ', False), ('true', False), (',   if  ', False), ('S × 100 ≥ T', False)], '3.15')

# 3.5.3 Device binding
h3('3.5.3  Device Binding')
body(
    'Each device is assigned a unique, persistent fingerprint U derived from its '
    'hardware and platform attributes. On first authentication the fingerprint is '
    'bound to the user account; thereafter, every sensitive request must present '
    'the same identifier. The binding rule is:'
)
equation([('Allow', True), (' = ', False), ('true', False), (',   if  ', False), ('U', False), ('req', False), (' = ', False), ('U', False), ('bound', False), ('  ∧  revoked_at = NULL', False)], '3.16')
where([
    ('U', 'req = the device identifier presented on the request;'),
    ('U', 'bound = the identifier stored against the account at first login.'),
])
body(
    'A login from a second device (Uᵣₑₒ ≠ Uᵇₒᵤₙᵈ) is rejected until an '
    'authorised reset revokes the previous binding, directly mitigating proxy '
    'attendance by a different physical device.'
)

# 3.5.4 Attendance timing & status
h3('3.5.4  Attendance Time Window and Status Classification')
body(
    'A session is only active within its scheduled window [tₛ, tₑ]. A check-in '
    'attempt at time t is admitted on the temporal axis only if:'
)
equation([('tₛ ≤ t < tₑ  ∧  status = active', False)], '3.17')
body('The lateness of an admitted check-in is measured as:')
equation([('m = (t − tₛ)  [minutes]', False)], '3.18')
body('and the attendance status is classified against a tolerance τ (the late threshold):')
equation([('Status', True), (' = ', False), ('PRESENT', False), (',  if  ', False), ('m ≤ τ;    ', False), ('LATE', False), (',  if  ', False), ('m > τ', False)], '3.19')
where([
    ('tₛ, tₑ', '= the scheduled start and end times of the session;'),
    ('m', '= minutes elapsed from the start time to the check-in instant;'),
    ('τ', '= the configured late-after threshold (minutes).'),
])

# 3.5.5 Continuous presence
h3('3.5.5  Continuous Presence Verification')
body(
    'To confirm that a student remains in the venue for the duration of the '
    'lecture — not merely at sign-in — the backend issues randomised presence '
    'checks during an active session. A new check is dispatched to checked-in '
    'students only when the elapsed time since the previous check exceeds the '
    'configured interval I:'
)
equation([('Trigger', True), (' = ', False), ('true', False), (',   if  ', False), ('(t', False), ('now', False), (' − t', False), ('last', False), (') ≥ I', False)], '3.20')
body(
    'Each presence check expires after a fixed response window (5 minutes); a '
    'student who fails to confirm presence from inside the venue accrues a missed '
    'check. The integrity of a record is captured by the present-throughout flag:'
)
equation([('PresentThroughout', True), (' = ', False), ('true', False), ('  ⇔  ', False), ('M = 0', False)], '3.21')
where([
    ('M', '= the number of missed presence checks for the record;'),
    ('I', '= the (randomised) presence-check interval (minutes).'),
])
body(
    'Finally, a student’s overall attendance rate over a course is computed for '
    'reporting and analytics as:'
)
equation([('Attendance Rate = (Aₐ / Aₜ) × 100  %', False)], '3.22')
where([
    ('Aₐ', '= the number of sessions attended (present or late);'),
    ('Aₜ', '= the total number of sessions held for the course.'),
])

# 3.5.6 Composite decision
h3('3.5.6  Composite Check-in Decision')
body(
    'The final attendance decision is the logical conjunction (AND) of all '
    'verification gates. Attendance is recorded only when every condition is '
    'satisfied simultaneously; the failure of any single gate produces a specific '
    'rejection reason:'
)
equation([
    ('Granted', True),
    (' = ', False),
    ('B ∧ G ∧ F ∧ W', False),
], '3.23')
where([
    ('B', '= device is bound and valid (Eq. 3.16);'),
    ('G', '= position lies inside the venue geofence (Eq. 3.6 / 3.7);'),
    ('F', '= face is verified against the enrolled reference (Eq. 3.13 / 3.15);'),
    ('W', '= the attempt falls within the active time window (Eq. 3.17).'),
])
body(
    'Only when G = B = F = W = true is an attendance record created and the '
    'student’s status classified by Eq. 3.19. This layered, multi-factor design '
    'simultaneously addresses location fraud, identity fraud, device-based proxy '
    'attendance, and out-of-window sign-ins.'
)

# ---- 3.6 Database design ----
h2('3.6  Database Design')
body(
    'A normalised relational schema was designed to eliminate redundancy and '
    'preserve referential integrity. The principal entities and their roles are '
    'summarised in Table 3.1.'
)
simple_table(
    ['Entity', 'Description', 'Key Relationships'],
    [
        ['users', 'Students, lecturers, HOCs, and administrators', 'has-many devices, enrollments, records'],
        ['institutions', 'Academic institutions hosting courses', 'has-many users, courses'],
        ['courses', 'Lecture classes with venue and schedule', 'belongs-to lecturer; has-one geofence'],
        ['course_geofences', 'Circular/polygon venue boundary', 'belongs-to course'],
        ['attendance_sessions', 'Scheduled/active lecture sittings', 'belongs-to course; has-many records'],
        ['attendance_records', 'Per-student check-in/out and status', 'belongs-to session, user, device'],
        ['face_profiles', 'Enrolled reference face descriptor', 'belongs-to user'],
        ['devices', 'Bound device fingerprints', 'belongs-to user'],
        ['presence_checks', 'Randomised in-lecture verifications', 'belongs-to session'],
    ],
    'Table 3.1: Principal database entities and relationships'
)

# ---- 3.7 Security ----
h2('3.7  System Security')
bullet('Transport security: all client–server traffic is encrypted with SSL/TLS over HTTPS.')
bullet('Authentication: stateless bearer tokens (Laravel Sanctum) authenticate every protected request.')
bullet('Password protection: user passwords are stored as one-way bcrypt hashes and are never recoverable in plaintext.')
bullet('Device binding: a unique device fingerprint is required on every sensitive request (Eq. 3.16).')
bullet('Multi-factor attendance: location, biometric, device, and time gates must all pass (Eq. 3.23).')

# ---- 3.8 System testing ----
h2('3.8  Method of System Evaluation')
body(
    'The implemented system was evaluated by functional and performance testing '
    'along the same path a genuine attendance follows: account registration and '
    'device binding, geofence (location) verification, face (identity) '
    'verification, the composite check-in decision under both legitimate and '
    'fraudulent conditions, the timed attendance window with randomised presence '
    'checks, and notification delivery with dashboard reporting. Positioning '
    'accuracy was assessed by comparing measured distances against ground-truth '
    'distances at known offsets from the venue centre; face verification was '
    'assessed using genuine, impostor, and spoof attempts; and platform stability '
    'was assessed under continuous operation and concurrent multi-user check-in. '
    'The results of these evaluations are presented and discussed in Chapter Four.'
)

doc.save('GeoTrack_Chapter_3_Methodology.docx')
print('Saved GeoTrack_Chapter_3_Methodology.docx')
