# -*- coding: utf-8 -*-
"""Turn the exported GeoTrack report into a submission-ready .docx.

Adds the four things requested:
  1. Abstract
  2. Table of Contents   (Word TOC field, headings 1-3)
  3. List of Figures      (Word TOC field on the FigureCaption style)
  4. List of Tables       (Word TOC field on the TableCaption style)
  5. Page numbers         (roman for front matter, arabic for the body)

Before that it REPAIRS the styling, because the Google-Docs export mislabelled
most body text as "Heading 2" and left Chapters 4-5 with no real heading
structure. Without this the TOC would list every paragraph.

Fields populate when opened in Word: select all (Ctrl+A) then press F9, or
right-click each field and choose "Update Field".
"""
import copy
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = "GeoTrack_Report_export.docx"
OUT = "GeoTrack_Report_FINAL.docx"

CENTER = WD_ALIGN_PARAGRAPH.CENTER

doc = Document(SRC)

# ── regexes ─────────────────────────────────────────────────────────────────
CH_RE = re.compile(r"^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT)\b", re.I)
SUB_RE = re.compile(r"^\d+\.\d+\.\d+")            # 3.3.1
SEC_RE = re.compile(r"^\d+\.\d+(?!\d)")           # 3.10  (checked after SUB)
FIG_RE = re.compile(r"^(Figure|Fig\.)\s*\d+(\.\d+)?\s*:", re.I)
TAB_RE = re.compile(r"^Table\s*\d+(\.\d+)?\s*:", re.I)


# ── style helpers ───────────────────────────────────────────────────────────
def style_names():
    return {s.name for s in doc.styles}


def ensure_para_style(name, *, italic=False, bold=False, size=12, caps=False, center=False):
    if name in style_names():
        return doc.styles[name]
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    st.font.italic = italic
    st.font.bold = bold
    st.font.size = Pt(size)
    if center:
        st.paragraph_format.alignment = CENTER
    return st


def set_style(p, name):
    p.style = doc.styles[name]


def set_text(p, text):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.add_run(text)


# ── Word field helper (TOC / PAGE) ──────────────────────────────────────────
def add_field(paragraph, instr, cached="Update field: select all then F9"):
    def run_with(child):
        r = OxmlElement("w:r")
        r.append(child)
        paragraph._p.append(r)

    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); run_with(fb)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; run_with(it)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); run_with(fs)
    tt = OxmlElement("w:t"); tt.set(qn("xml:space"), "preserve"); tt.text = cached; run_with(tt)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); run_with(fe)


# ── custom styles ───────────────────────────────────────────────────────────
ensure_para_style("FigureCaption", italic=True, size=11, center=True)
ensure_para_style("TableCaption", italic=True, size=11, center=True)
ensure_para_style("FrontMatterHeading", bold=True, size=14, center=True)

# ============================================================================
# 1. REPAIR STYLES
# ============================================================================
plist = list(doc.paragraphs)
to_remove = []
chapter_one = None
n = len(plist)
k = 0
while k < n:
    p = plist[k]
    t = p.text.strip()
    if not t:
        k += 1
        continue

    if CH_RE.match(t):
        # find the chapter-name line that follows (e.g. INTRODUCTION)
        m = k + 1
        while m < n and not plist[m].text.strip():
            m += 1
        name = plist[m].text.strip() if m < n else ""
        merged = t
        if (name and not re.match(r"^\d", name) and not CH_RE.match(name)
                and len(name.split()) <= 8 and name.upper() == name):
            merged = f"{t}: {name}"
            to_remove.append(plist[m])
            advance = m + 1
        else:
            advance = k + 1
        set_text(p, merged)
        set_style(p, "Heading 1")
        p.alignment = CENTER
        if chapter_one is None and re.search(r"CHAPTER\s+ONE", merged, re.I):
            chapter_one = p
        k = advance
        continue

    if SUB_RE.match(t):
        set_style(p, "Heading 3"); k += 1; continue
    if SEC_RE.match(t):
        set_style(p, "Heading 2"); k += 1; continue
    if FIG_RE.match(t):
        set_style(p, "FigureCaption"); p.alignment = CENTER; k += 1; continue
    if TAB_RE.match(t):
        set_style(p, "TableCaption"); p.alignment = CENTER; k += 1; continue
    if t.upper() == "REFERENCES":
        set_style(p, "Heading 1"); p.alignment = CENTER; k += 1; continue

    # demote anything still wrongly tagged as a heading
    sn = p.style.name
    if sn == "Heading 4":
        set_style(p, "Heading 3")
    elif sn in ("Heading 2", "Heading 5", "Heading 6"):
        set_style(p, "Normal")
    k += 1

for p in to_remove:
    p._element.getparent().remove(p._element)

if chapter_one is None:
    raise SystemExit("Could not locate CHAPTER ONE anchor.")

# ============================================================================
# 2. ABSTRACT
# ============================================================================
ABSTRACT = (
    "Accurate attendance monitoring remains a persistent challenge in academic "
    "institutions, where conventional manual and basic digital methods are "
    "vulnerable to impersonation (proxy attendance), location fraud, and the "
    "inability to confirm a student's continued presence throughout a lecture. "
    "This project presents the design and implementation of GeoTrack, a "
    "geo-fenced mobile attendance monitoring system that integrates four "
    "independent verification mechanisms - device binding, GPS-based geofencing, "
    "facial verification, and a lecturer-controlled timed attendance window with "
    "randomised presence checks - so that attendance is recorded only when a "
    "genuine, physically present, and correctly identified student checks in from "
    "a registered device.\n"
    "The system was developed as a three-tier client-server application "
    "comprising a cross-platform React Native mobile application, a Laravel "
    "RESTful backend, and a relational database, communicating securely over "
    "HTTPS. A student's position is determined by averaging multiple "
    "high-accuracy GNSS samples and evaluated against circular or polygonal "
    "venue boundaries using the Haversine and ray-casting algorithms, while "
    "identity is confirmed through facial verification against a configurable "
    "similarity threshold.\n"
    "The implemented system was evaluated through functional and performance "
    "testing covering positioning accuracy, face-verification reliability, "
    "device-binding enforcement, and the composite check-in decision under both "
    "legitimate and fraudulent conditions. Results showed accurate geofence "
    "admission outside the GPS-margin zone, reliable rejection of impersonation "
    "and location-spoofing attempts, correct on-time and late classification, "
    "prompt push and email notifications, and stable operation under continuous "
    "and concurrent use. GeoTrack therefore demonstrates a secure, "
    "cost-effective, and engineering-driven solution that eliminates common "
    "attendance fraud while providing real-time monitoring and analytics for "
    "lecturers and administrators."
)


def insert_heading(anchor, text):
    p = anchor.insert_paragraph_before(text, style="FrontMatterHeading")
    p.alignment = CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(12)
    return p


def insert_body(anchor, text):
    p = anchor.insert_paragraph_before(text, style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def insert_field(anchor, instr):
    p = anchor.insert_paragraph_before("", style="Normal")
    add_field(p, instr)
    return p


# Order (each inserted directly above CHAPTER ONE, so call order == page order):
insert_heading(chapter_one, "ABSTRACT")
for para in ABSTRACT.split("\n"):
    insert_body(chapter_one, para)

insert_heading(chapter_one, "TABLE OF CONTENTS")
insert_field(chapter_one, 'TOC \\o "1-3" \\h \\z \\u')

insert_heading(chapter_one, "LIST OF FIGURES")
insert_field(chapter_one, 'TOC \\h \\z \\t "FigureCaption,1"')

insert_heading(chapter_one, "LIST OF TABLES")
insert_field(chapter_one, 'TOC \\h \\z \\t "TableCaption,1"')

# ============================================================================
# 3. SECTION BREAK + PAGE NUMBERS  (front matter = roman, body = arabic)
# ============================================================================
break_para = chapter_one.insert_paragraph_before("", style="Normal")
body_sectPr = doc.sections[-1]._sectPr
fm_sectPr = copy.deepcopy(body_sectPr)
break_para._p.get_or_add_pPr().append(fm_sectPr)


def set_pgnum(section, fmt, start=None):
    sectPr = section._sectPr
    el = sectPr.find(qn("w:pgNumType"))
    if el is None:
        el = OxmlElement("w:pgNumType")
        cols = sectPr.find(qn("w:cols"))
        if cols is not None:
            cols.addprevious(el)
        else:
            sectPr.append(el)
    el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))


def footer_page_number(section, different_first=False):
    section.different_first_page_header_footer = different_first
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = CENTER
    add_field(p, "PAGE", cached="")
    if different_first:
        ff = section.first_page_footer
        ff.is_linked_to_previous = False
        # leave the title-page footer empty (no number printed there)


# python-docx now sees two sections: [0] front matter, [1] body.
front, body = doc.sections[0], doc.sections[-1]
set_pgnum(front, "lowerRoman")
set_pgnum(body, "decimal", start=1)
footer_page_number(front, different_first=True)  # hide number on title page
footer_page_number(body, different_first=False)

doc.save(OUT)
print(f"Saved {OUT}")
print(f"Sections now: {len(doc.sections)}")
